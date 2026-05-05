import os
import csv
import io
import random
import json
from datetime import datetime, date, timedelta
from dotenv import load_dotenv
from google import genai
from werkzeug.utils import secure_filename
from flask import Flask, render_template, request, redirect, url_for, flash, Response, jsonify, send_file, session
from docx import Document
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import func, inspect
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from flask_mail import Mail, Message

# --- BULLETPROOF ENV LOADER ---
# This forces Python to find the .env file exactly where app.py lives
basedir = os.path.abspath(os.path.dirname(__file__))
load_dotenv(os.path.join(basedir, '.env'), override=True)

# --- GEMINI AI (Optional) ---
api_key = os.environ.get("GEMINI_API_KEY", "").strip()
client = None
if api_key and not api_key.startswith("your_"):
    try:
        client = genai.Client(api_key=api_key)
        print(f"\n--- Gemini AI: Enabled (key starts with: {api_key[:6]}) ---\n")
    except Exception as e:
        print(f"\n--- Gemini AI: Failed to init ({e}) — AI features disabled ---\n")
else:
    print("\n--- Gemini AI: No valid key found — AI features disabled ---\n")

# --- Ensure upload directory exists ---
os.makedirs('static/uploads', exist_ok=True)

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///icar_portfolio.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-fallback-key-change-in-production')

# --- EMAIL CONFIGURATION ---
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = os.environ.get('MAIL_USERNAME', '')
app.config['MAIL_PASSWORD'] = os.environ.get('MAIL_PASSWORD', '')
mail = Mail(app)

db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'

# Register fromjson Jinja2 filter so templates can parse JSON strings
import json as _json
app.jinja_env.filters['fromjson'] = lambda s: _json.loads(s) if s else []

# --- SERVER-SIDE OTP STORE (fixes cookie leak vulnerability) ---
_otp_store = {}

# --- DATABASE MODELS ---

class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(150), unique=True, nullable=False)
    email = db.Column(db.String(150), unique=True, nullable=True)
    password = db.Column(db.String(150), nullable=False)
    role = db.Column(db.String(50), nullable=False)
    is_approved = db.Column(db.Boolean, default=False)

class Variety(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    ip_asset_id = db.Column(db.String(50), unique=True, nullable=False)
    status = db.Column(db.String(50), default='Filed')
    category = db.Column(db.String(50)) 
    year_of_release = db.Column(db.Integer)
    duration_days = db.Column(db.Integer) 
    ppvfr_registration = db.Column(db.String(10)) 
    agency_released_by = db.Column(db.String(100))
    yield_data = db.Column(db.String(100)) 
    recommended_zones = db.Column(db.Text) 
    states = db.Column(db.Text)
    grain_types = db.Column(db.Text) 
    pest_resistances = db.Column(db.Text) 
    disease_resistances = db.Column(db.Text)
    abiotic_stress = db.Column(db.Text)
    special_traits = db.Column(db.Text)
    mou_copy_filename = db.Column(db.String(200))

class VarietyLicense(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    ip_asset_id = db.Column(db.String(50), nullable=False)
    category = db.Column(db.String(50))
    company_licensed = db.Column(db.String(100))
    date_licensed = db.Column(db.Date)
    license_fee = db.Column(db.Float)
    royalty_received = db.Column(db.Float)
    date_royalty_received = db.Column(db.Date)
    mou_copy_filename = db.Column(db.String(200)) # <-- ADD THIS LINE
    
class PatentProduct(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    ip_asset_id = db.Column(db.String(50), unique=True, nullable=False)
    status = db.Column(db.String(50), default='Filed')
    patent_name = db.Column(db.String(200))
    date_filed = db.Column(db.Date)
    date_received = db.Column(db.Date)
    valid_up_to = db.Column(db.Date)
    company_licensed = db.Column(db.String(100))
    date_licensed = db.Column(db.Date)
    mou_date = db.Column(db.Date)
    mou_copy_filename = db.Column(db.String(200))
    date_granted = db.Column(db.Date)
    license_fee = db.Column(db.Float)
    license_fee_date = db.Column(db.Date) 
    royalty_received = db.Column(db.Float)
    date_royalty_received = db.Column(db.Date) 
    
class PatentProcess(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    ip_asset_id = db.Column(db.String(50), unique=True, nullable=False)
    status = db.Column(db.String(50), default='Filed')
    process_name = db.Column(db.String(200))
    date_filed = db.Column(db.Date)
    date_received = db.Column(db.Date)
    valid_up_to = db.Column(db.Date)
    company_licensed = db.Column(db.String(100))
    date_licensed = db.Column(db.Date)
    mou_date = db.Column(db.Date)
    mou_copy_filename = db.Column(db.String(200))
    date_granted = db.Column(db.Date)
    license_fee = db.Column(db.Float)
    license_fee_date = db.Column(db.Date) 
    royalty_received = db.Column(db.Float)
    date_royalty_received = db.Column(db.Date) 

class PatentDesign(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    ip_asset_id = db.Column(db.String(50), unique=True, nullable=False)
    status = db.Column(db.String(50), default='Filed')
    design_name = db.Column(db.String(200)) 
    date_filed = db.Column(db.Date)
    date_received = db.Column(db.Date)
    valid_up_to = db.Column(db.Date)
    company_licensed = db.Column(db.String(100))
    date_licensed = db.Column(db.Date)
    mou_date = db.Column(db.Date)
    mou_copy_filename = db.Column(db.String(200))
    date_granted = db.Column(db.Date)
    license_fee = db.Column(db.Float)
    license_fee_date = db.Column(db.Date) 
    royalty_received = db.Column(db.Float)
    date_royalty_received = db.Column(db.Date)
    
class Copyright(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    ip_asset_id = db.Column(db.String(50), unique=True, nullable=False)
    status = db.Column(db.String(50), default='Filed')
    article_name = db.Column(db.String(200))
    date_filed = db.Column(db.Date)
    date_received = db.Column(db.Date)
    valid_up_to = db.Column(db.Date)
    company_licensed = db.Column(db.String(100))
    date_licensed = db.Column(db.Date)
    license_fee = db.Column(db.Float)
    mou_copy_filename = db.Column(db.String(200))
    royalty_received = db.Column(db.Float)
    date_royalty_received = db.Column(db.Date)
    
class Trademark(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    ip_asset_id = db.Column(db.String(50), unique=True, nullable=False)
    status = db.Column(db.String(50), default='Filed')
    asset_type = db.Column(db.String(50)) 
    asset_name = db.Column(db.String(200))
    date_filed = db.Column(db.Date)
    date_received = db.Column(db.Date)
    valid_up_to = db.Column(db.Date)
    company_licensed = db.Column(db.String(100))
    date_licensed = db.Column(db.Date)
    license_fee = db.Column(db.Float)
    mou_copy_filename = db.Column(db.String(200))
    royalty_received = db.Column(db.Float)
    date_royalty_received = db.Column(db.Date)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

def calculate_patent_deadlines(filing_date):
    if not filing_date: return None
    try:
        return filing_date.replace(year=filing_date.year + 20)
    except ValueError:
        return filing_date.replace(year=filing_date.year + 20, day=28)

@app.route('/setup')
def setup():
    db.create_all()
    msg = ""
    if not User.query.filter_by(username='admin').first():
        admin_email = os.environ.get('ADMIN_EMAIL') or os.environ.get('MAIL_USERNAME', 'admin@example.com')
        admin_password = os.environ.get('ADMIN_PASSWORD', 'admin123')
        admin = User(username='admin', email=admin_email, password=generate_password_hash(admin_password), role='Admin', is_approved=True)
        db.session.add(admin)
        msg += "Admin created. "
    db.session.commit()
    return f"{msg} Database ready! Go to <a href='/login'>/login</a>"

@app.route('/', methods=['GET', 'POST'])
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        user = User.query.filter_by(username=request.form.get('username')).first()
        if user and check_password_hash(user.password, request.form.get('password')):
            if not user.is_approved:
                flash('Your account is pending Admin approval. Please check back later.')
                return redirect(url_for('login'))
            
            otp = str(random.randint(100000, 999999))
            session['pending_user_id'] = user.id
            _otp_store[user.id] = otp  # Stored server-side, NOT in cookie
            print(f"--- SECURITY LOG: OTP FOR {user.username} IS {otp} ---")
            try:
                msg = Message("Your ICAR Security PIN", sender=os.environ.get('MAIL_USERNAME', 'admin@example.com'), recipients=[user.email])
                msg.body = f"Your one-time security PIN is: {otp}\nDo not share this code with anyone."
                mail.send(msg)
                flash('A 6-digit security PIN has been sent to your email.')
            except Exception as e:
                print(e)
                flash('Check terminal for OTP (Email send failed).')
                
            return redirect(url_for('two_factor'))
        
        flash('Invalid Credentials')
    return render_template('login.html')

@app.route('/2fa', methods=['GET', 'POST'])
def two_factor():
    if 'pending_user_id' not in session:
        return redirect(url_for('login'))

    if request.method == 'POST':
        # Clean up the input and ensure string comparison
        user_pin = request.form.get('otp', '').strip()
        expected = str(_otp_store.get(session.get('pending_user_id', -1), '')).strip()

        print(f">>> DEBUG: Typed [{user_pin}], Expected [{expected}]")

        if user_pin and user_pin == expected:
            user_id = session['pending_user_id']
            user = User.query.get(user_id)
            login_user(user)
            _otp_store.pop(user_id, None)  # Clean up server-side OTP
            session.pop('pending_user_id', None)
            return redirect(url_for('dashboard'))
        else:
            flash(f"Invalid PIN. (Typed: {user_pin})")

    return render_template('verify_otp.html')

@app.route('/dashboard')
@login_required 
def dashboard():
    total_v = Variety.query.count()
    total_p = PatentProduct.query.count() + PatentProcess.query.count()
    total_c = Copyright.query.count()
    total_d = PatentDesign.query.count()
    total_tm = Trademark.query.count()
    
    total_assets = total_v + total_p + total_c + total_d + total_tm
    total_brands = total_c + total_tm # Combined count for Copyrights + Trademarks

    r_prod = (db.session.query(func.sum(PatentProduct.license_fee)).scalar() or 0) + (db.session.query(func.sum(PatentProduct.royalty_received)).scalar() or 0)
    r_proc = (db.session.query(func.sum(PatentProcess.license_fee)).scalar() or 0) + (db.session.query(func.sum(PatentProcess.royalty_received)).scalar() or 0)
    r_des = (db.session.query(func.sum(PatentDesign.license_fee)).scalar() or 0) + (db.session.query(func.sum(PatentDesign.royalty_received)).scalar() or 0)
    r_copy = (db.session.query(func.sum(Copyright.license_fee)).scalar() or 0) + (db.session.query(func.sum(Copyright.royalty_received)).scalar() or 0)
    r_var = (db.session.query(func.sum(VarietyLicense.license_fee)).scalar() or 0) + (db.session.query(func.sum(VarietyLicense.royalty_received)).scalar() or 0)
    r_tm = (db.session.query(func.sum(Trademark.license_fee)).scalar() or 0) + (db.session.query(func.sum(Trademark.royalty_received)).scalar() or 0)

    revenue = r_prod + r_proc + r_des + r_copy + r_var + r_tm
    
    today = date.today()
    expiring_soon = PatentProduct.query.filter(
        PatentProduct.valid_up_to >= today,
        PatentProduct.valid_up_to <= today + timedelta(days=30)
    ).all()

    status_counts = {'Filed': 0, 'Granted': 0, 'Licensed': 0}
    for model in [Variety, PatentProduct, PatentProcess, PatentDesign, Copyright, Trademark]: 
        status_counts['Filed'] += model.query.filter(model.status.ilike('%Filed%')).count()
        status_counts['Granted'] += model.query.filter(model.status.ilike('%Granted%')).count()
        status_counts['Licensed'] += model.query.filter(model.status.ilike('%Licensed%')).count()

    return render_template('dashboard.html', 
                           username=current_user.username, 
                           role=current_user.role,
                           total_varieties=total_v,
                           total_patents=total_p,
                           total_brands=total_brands,  # Passed combined total_brands
                           total_designs=total_d,
                           total_assets=total_assets,
                           revenue=revenue,
                           expiring_soon=expiring_soon,
                           status_counts=status_counts)
@app.route('/repository')
@login_required
def repository():
    return render_template('repository.html', role=current_user.role)

@app.route('/repository/<repo_type>')
@login_required
def repository_view(repo_type):
    data = {}
    status_counts = {'Filed': 0, 'Granted': 0, 'Licensed': 0}
    
    def clean_json(json_str):
        if not json_str or json_str == '[]': return "None"
        try:
            parsed = json.loads(json_str)
            if isinstance(parsed, list):
                if len(parsed) > 0 and isinstance(parsed[0], dict):
                    return ", ".join([f"{item.get('name', '')} ({item.get('level', '')})" for item in parsed])
                else:
                    return ", ".join(parsed)
            return str(parsed)
        except:
            return json_str

    if repo_type == 'varieties':
        items = Variety.query.all()
        for i in items:
            i.clean_zones = clean_json(i.recommended_zones)
            i.clean_states = clean_json(i.states)
            i.clean_grains = clean_json(i.grain_types) 
            i.clean_pests = clean_json(i.pest_resistances)
            i.clean_diseases = clean_json(i.disease_resistances)
            i.clean_abiotic = clean_json(i.abiotic_stress)
        data['items'] = items
        status_counts['Filed'] = Variety.query.filter(Variety.status.ilike('%Filed%')).count()
        status_counts['Granted'] = Variety.query.filter(Variety.status.ilike('%Granted%')).count()
        status_counts['Licensed'] = Variety.query.filter(Variety.status.ilike('%Licensed%')).count()
        
    elif repo_type == 'patents':
        data['prod'] = PatentProduct.query.all()
        data['proc'] = PatentProcess.query.all()
        data['des'] = PatentDesign.query.all()
        for model in [PatentProduct, PatentProcess, PatentDesign]:
            status_counts['Filed'] += model.query.filter(model.status.ilike('%Filed%')).count()
            status_counts['Granted'] += model.query.filter(model.status.ilike('%Granted%')).count()
            status_counts['Licensed'] += model.query.filter(model.status.ilike('%Licensed%')).count()
        
    elif repo_type == 'brands':
        data['copy'] = Copyright.query.all()
        data['tm'] = Trademark.query.all()
        for model in [Copyright, Trademark]:
            status_counts['Filed'] += model.query.filter(model.status.ilike('%Filed%')).count()
            status_counts['Granted'] += model.query.filter(model.status.ilike('%Granted%')).count()
            status_counts['Licensed'] += model.query.filter(model.status.ilike('%Licensed%')).count()
        
    elif repo_type == 'licenses':
        data['vl'] = VarietyLicense.query.all()
        data['prod'] = PatentProduct.query.filter_by(status='Licensed').all()
        data['proc'] = PatentProcess.query.filter_by(status='Licensed').all()
        data['des'] = PatentDesign.query.filter_by(status='Licensed').all()
        data['copy'] = Copyright.query.filter_by(status='Licensed').all()
        data['tm'] = Trademark.query.filter_by(status='Licensed').all()
        status_counts['Licensed'] = len(data['vl']) + len(data['prod']) + len(data['proc']) + len(data['des']) + len(data['copy']) + len(data['tm'])
        
    return render_template('repository_details.html', repo_type=repo_type, data=data, status_counts=status_counts, role=current_user.role)

@app.route('/add_variety', methods=['GET', 'POST'])
@login_required
def add_variety():
    if current_user.role != 'Admin': return redirect(url_for('dashboard'))
    if request.method == 'POST':
        asset_id = request.form.get('ip_asset_id')
        if Variety.query.filter_by(ip_asset_id=asset_id).first():
            flash(f"Error: Asset ID '{asset_id}' already exists!")
            return redirect(url_for('add_variety'))

        # --- NEW FILE UPLOAD LOGIC ---
        file = request.files.get('mou_file')
        fname = secure_filename(file.filename) if file and file.filename != '' else None
        if fname:
            fname = f"VAR_{asset_id}_{fname}"
            import os
            # Saves to your existing static/uploads folder
            file.save(os.path.join('static/uploads', fname)) 
        # -----------------------------

        zones = request.form.getlist('zones[]')
        states = request.form.getlist('states[]')
        grains = request.form.getlist('grain_types[]')
        
        pests = []
        p_names = request.form.getlist('pest_n[]')
        p_levels = request.form.getlist('pest_l[]')
        for n, l in zip(p_names, p_levels):
            if n: pests.append({"name": n, "level": l})

        diseases = []
        d_names = request.form.getlist('disease_n[]')
        d_levels = request.form.getlist('disease_l[]')
        for n, l in zip(d_names, d_levels):
            if n: diseases.append({"name": n, "level": l})

        abiotic = []
        a_names = request.form.getlist('abiotic_n[]')
        a_levels = request.form.getlist('abiotic_l[]')
        for n, l in zip(a_names, a_levels):
            if n: abiotic.append({"name": n, "level": l})

        v = Variety(
            ip_asset_id=asset_id,
            status=request.form.get('status').strip(),
            category=request.form.get('category'),
            year_of_release=request.form.get('year_of_release'),
            duration_days=request.form.get('duration_days'),
            ppvfr_registration=request.form.get('ppvfr'),
            agency_released_by=request.form.get('agency_released_by'),
            yield_data=request.form.get('yield_data'),
            recommended_zones=json.dumps(zones),
            states=json.dumps(states),
            grain_types=json.dumps(grains), 
            pest_resistances=json.dumps(pests),
            disease_resistances=json.dumps(diseases),
            abiotic_stress=json.dumps(abiotic),
            special_traits=request.form.get('special_traits'),
            mou_copy_filename=fname  # <-- ADDED THIS LINE
        )
        db.session.add(v)
        db.session.commit()
        return redirect(url_for('dashboard'))
    return render_template('add_variety.html')

@app.route('/record_license', methods=['GET', 'POST'])
@login_required
def record_license():
    if current_user.role != 'Admin': return redirect(url_for('dashboard'))
    if request.method == 'POST':
        file = request.files.get('mou_file')
        fname = secure_filename(file.filename) if file and file.filename != '' else None
        if fname: file.save(os.path.join('static/uploads', fname))

        def p_date(d): return datetime.strptime(d, '%Y-%m-%d').date() if d else None
        
        asset_id = request.form.get('ip_asset_id')
        category = request.form.get('category')
        company = request.form.get('company_licensed')
        
        date_lic = p_date(request.form.get('date_licensed'))
        date_fee_rec = p_date(request.form.get('license_fee_date'))
        date_roy_rec = p_date(request.form.get('date_royalty_received'))
        
        l_fee = request.form.get('license_fee') or 0.0
        r_rec = request.form.get('royalty_received') or 0.0

        exists = False
        if category in ['Variety', 'Hybrid']: exists = Variety.query.filter_by(ip_asset_id=asset_id).first()
        elif category == 'Product': exists = PatentProduct.query.filter_by(ip_asset_id=asset_id).first()
        elif category == 'Process': exists = PatentProcess.query.filter_by(ip_asset_id=asset_id).first()
        elif category == 'Design': exists = PatentDesign.query.filter_by(ip_asset_id=asset_id).first()
        elif category == 'Copyright': exists = Copyright.query.filter_by(ip_asset_id=asset_id).first()
        elif category in ['Trademark', 'Logo']: exists = Trademark.query.filter_by(ip_asset_id=asset_id).first()

        if not exists:
            flash(f"Error: Asset ID '{asset_id}' does not exist in category '{category}'.")
            return redirect(url_for('record_license'))

        if category in ['Variety', 'Hybrid']:
            vl = VarietyLicense(
                ip_asset_id=asset_id, category=category, company_licensed=company,
                date_licensed=date_lic, license_fee=l_fee, royalty_received=r_rec,
                date_royalty_received=date_roy_rec,
                mou_copy_filename=fname # <-- ADD THIS HERE
            )
            db.session.add(vl)
            v = Variety.query.filter_by(ip_asset_id=asset_id).first()
            if v: v.status = 'Licensed'

        else:
            item = None
            if category == 'Product': item = PatentProduct.query.filter_by(ip_asset_id=asset_id).first()
            elif category == 'Process': item = PatentProcess.query.filter_by(ip_asset_id=asset_id).first()
            elif category == 'Design': item = PatentDesign.query.filter_by(ip_asset_id=asset_id).first()
            elif category == 'Copyright': item = Copyright.query.filter_by(ip_asset_id=asset_id).first()
            elif category in ['Trademark', 'Logo']: item = Trademark.query.filter_by(ip_asset_id=asset_id).first()

            if item:
                item.status = 'Licensed'
                item.company_licensed = company
                item.date_licensed = date_lic
                item.license_fee = l_fee
                item.license_fee_date = date_fee_rec
                item.royalty_received = r_rec
                item.date_royalty_received = date_roy_rec
                if hasattr(item, 'mou_copy_filename'): 
                    item.mou_copy_filename = fname

        db.session.commit()
        return redirect(url_for('dashboard'))
    return render_template('record_license.html')

@app.route('/add_patent', methods=['GET', 'POST'])
@login_required
def add_patent():
    if current_user.role != 'Admin': return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        asset_id = request.form.get('ip_asset_id')
        if PatentProduct.query.filter_by(ip_asset_id=asset_id).first() or PatentProcess.query.filter_by(ip_asset_id=asset_id).first() or PatentDesign.query.filter_by(ip_asset_id=asset_id).first():
            flash(f"Error: Asset ID '{asset_id}' already exists!")
            return redirect(url_for('add_patent'))

        file = request.files.get('mou_file')
        fname = secure_filename(file.filename) if file and file.filename != '' else None
        if fname: file.save(os.path.join('static/uploads', fname))

        def p_date(d): return datetime.strptime(d, '%Y-%m-%d').date() if d else None
        filing_date = p_date(request.form.get('date_filed'))
        auto_valid_up_to = calculate_patent_deadlines(filing_date)

        data = {
            'ip_asset_id': asset_id,
            'status': request.form.get('status', 'Filed').strip(),
            'date_filed': filing_date,
            'date_granted': p_date(request.form.get('date_granted')),
            'valid_up_to': auto_valid_up_to,
            'company_licensed': request.form.get('company_licensed'),
            'mou_date': p_date(request.form.get('mou_date')),
            'mou_copy_filename': fname,
            'license_fee': request.form.get('license_fee') or 0.0,
            'license_fee_date': p_date(request.form.get('license_fee_date')),
            'royalty_received': request.form.get('royalty_received') or 0.0,
            'date_royalty_received': p_date(request.form.get('date_royalty_received')),
        }
        
        ptype = request.form.get('patent_type')
        title = request.form.get('patent_name')

        if ptype == 'Product':
            new = PatentProduct(patent_name=title, **data)
        elif ptype == 'Process':
            new = PatentProcess(process_name=title, **data)
        else:
            new = PatentDesign(design_name=title, **data)

        db.session.add(new)
        db.session.commit()
        return redirect(url_for('dashboard'))
        
    return render_template('add_patent.html')

@app.route('/add_ip', methods=['GET', 'POST'])
@login_required
def add_ip():
    if current_user.role != 'Admin': return redirect(url_for('dashboard'))
    if request.method == 'POST':
        asset_id = request.form.get('ip_asset_id')
        if Copyright.query.filter_by(ip_asset_id=asset_id).first() or Trademark.query.filter_by(ip_asset_id=asset_id).first():
            flash(f"Error: Asset ID '{asset_id}' already exists!")
            return redirect(url_for('add_ip'))

        file = request.files.get('mou_file')
        fname = secure_filename(file.filename) if file and file.filename != '' else None
        if fname: file.save(os.path.join('static/uploads', fname))

        def p_date(d): return datetime.strptime(d, '%Y-%m-%d').date() if d else None

        data = {
            'ip_asset_id': asset_id,
            'status': request.form.get('status', 'Filed').strip(),
            'company_licensed': request.form.get('company_licensed'),
            'date_licensed': p_date(request.form.get('date_licensed')),
            'license_fee': request.form.get('license_fee') or 0.0,
            'royalty_received': request.form.get('royalty_received') or 0.0,
            'date_royalty_received': p_date(request.form.get('date_royalty_received')),
            'mou_copy_filename': fname
        }
        
        category = request.form.get('asset_category')
        name = request.form.get('name')

        if category == 'Copyright':
            new = Copyright(article_name=name, **data)
        else:
            new = Trademark(asset_type=category, asset_name=name, **data)
            
        db.session.add(new)
        db.session.commit()
        return redirect(url_for('dashboard'))
        
    return render_template('add_ip.html')

@app.route('/delete/<asset_type>/<int:id>')
@login_required
def delete_asset(asset_type, id):
    if current_user.role != 'Admin': return redirect(url_for('repository'))
    item = None
    if asset_type == 'variety': item = Variety.query.get(id)
    elif asset_type == 'patent': item = PatentProduct.query.get(id)
    elif asset_type == 'copyright': item = Copyright.query.get(id)
    elif asset_type == 'design': item = PatentDesign.query.get(id)
    
    if item:
        db.session.delete(item)
        db.session.commit()
    return redirect(url_for('repository'))

@app.route('/export_csv/<repo_type>')
@login_required
def export_csv_detailed(repo_type):
    si = io.StringIO()
    cw = csv.writer(si)
    
    def clean_json(json_str):
        if not json_str or json_str == '[]': return "None"
        try:
            parsed = json.loads(json_str)
            if isinstance(parsed, list):
                if len(parsed) > 0 and isinstance(parsed[0], dict):
                    return ", ".join([f"{item.get('name', '')} ({item.get('level', '')})" for item in parsed])
                else:
                    return ", ".join(parsed)
            return str(parsed)
        except:
            return json_str

    if repo_type == 'varieties':
        filename = "Detailed_Varieties_Report.csv"
        cw.writerow(['ID', 'Cat', 'Status', 'Year', 'Duration', 'PPVFR', 'Notified By', 'Yield', 'Zones', 'States', 'Grain Types', 'Pests', 'Diseases', 'Abiotic Stress', 'Special Traits'])
        for v in Variety.query.all():
            cw.writerow([
                v.ip_asset_id, v.category, v.status, v.year_of_release, v.duration_days, 
                v.ppvfr_registration, v.agency_released_by, v.yield_data, 
                clean_json(v.recommended_zones), clean_json(v.states), clean_json(v.grain_types), 
                clean_json(v.pest_resistances), clean_json(v.disease_resistances), 
                clean_json(v.abiotic_stress), v.special_traits
            ])
            
    elif repo_type == 'patents':
        filename = "Detailed_Patents_Report.csv"
        cw.writerow(['ID', 'Type', 'Title', 'Status', 'Date Filed', 'Date Granted', 'Expiry', 'Company', 'License / MoU Date', 'Fee', 'Fee Date', 'Royalty', 'Royalty Date'])
        models = [(PatentProduct, 'Product Patent'), (PatentProcess, 'Process Patent'), (PatentDesign, 'Design Patent')]
        for m, label in models:
            for p in m.query.all():
                name = getattr(p, 'patent_name', getattr(p, 'process_name', getattr(p, 'design_name', '')))
                actual_date = p.date_licensed if p.date_licensed else p.mou_date
                cw.writerow([p.ip_asset_id, label, name, p.status, p.date_filed, p.date_granted, p.valid_up_to, p.company_licensed, actual_date, p.license_fee, p.license_fee_date, p.royalty_received, p.date_royalty_received])
            
    elif repo_type == 'brands':
        filename = "Detailed_Brands_Report.csv"
        cw.writerow(['ID', 'Type', 'Name', 'Status', 'Company', 'License / MoU Date', 'Fee', 'Royalty', 'Royalty Date'])
        for c in Copyright.query.all(): cw.writerow([c.ip_asset_id, 'Copyright', c.article_name, c.status, c.company_licensed, c.date_licensed, c.license_fee, c.royalty_received, c.date_royalty_received])
        for t in Trademark.query.all(): cw.writerow([t.ip_asset_id, t.asset_type, t.asset_name, t.status, t.company_licensed, t.date_licensed, t.license_fee, t.royalty_received, t.date_royalty_received])
            
    elif repo_type == 'licenses':
        filename = "Detailed_Commercialization_Report.csv"
        cw.writerow(['ID', 'Asset Type', 'Company Licensed', 'License / MoU Date', 'License Fee', 'Fee Received Date', 'Royalty', 'Royalty Received Date'])
        for v in VarietyLicense.query.all(): cw.writerow([v.ip_asset_id, v.category, v.company_licensed, v.date_licensed, v.license_fee, '-', v.royalty_received, v.date_royalty_received])
        
        for p in PatentProduct.query.filter_by(status='Licensed').all(): 
            actual_date = p.date_licensed if p.date_licensed else p.mou_date
            cw.writerow([p.ip_asset_id, 'Product', p.company_licensed, actual_date, p.license_fee, p.license_fee_date, p.royalty_received, p.date_royalty_received])
        for p in PatentProcess.query.filter_by(status='Licensed').all(): 
            actual_date = p.date_licensed if p.date_licensed else p.mou_date
            cw.writerow([p.ip_asset_id, 'Process', p.company_licensed, actual_date, p.license_fee, p.license_fee_date, p.royalty_received, p.date_royalty_received])
        for d in PatentDesign.query.filter_by(status='Licensed').all(): 
            actual_date = d.date_licensed if d.date_licensed else d.mou_date
            cw.writerow([d.ip_asset_id, 'Design', d.company_licensed, actual_date, d.license_fee, d.license_fee_date, d.royalty_received, d.date_royalty_received])
            
        for c in Copyright.query.filter_by(status='Licensed').all(): cw.writerow([c.ip_asset_id, 'Copyright', c.company_licensed, c.date_licensed, c.license_fee, '-', c.royalty_received, c.date_royalty_received])
        for t in Trademark.query.filter_by(status='Licensed').all(): cw.writerow([t.ip_asset_id, t.asset_type, t.company_licensed, t.date_licensed, t.license_fee, '-', t.royalty_received, t.date_royalty_received])

    output = Response(
        si.getvalue().encode('utf-8-sig'),  # BOM prefix so Excel opens correctly on Windows
        mimetype='text/csv; charset=utf-8'
    )
    output.headers["Content-Disposition"] = f"attachment; filename={filename}"
    return output

@app.route('/deadlines')
@login_required
def deadlines():
    try:
        schedule = []
        models = [PatentProduct, PatentProcess, PatentDesign, Copyright]
        
        for model in models:
            items = model.query.all()
            for i in items:
                if hasattr(i, 'patent_name') and i.patent_name: name = i.patent_name
                elif hasattr(i, 'process_name') and i.process_name: name = i.process_name
                elif hasattr(i, 'design_name') and i.design_name: name = i.design_name
                elif hasattr(i, 'article_name') and i.article_name: name = i.article_name
                else: name = "IP Asset"

                if hasattr(i, 'valid_up_to') and i.valid_up_to:
                    schedule.append({'title': f"{name} (EXPIRY)", 'start': str(i.valid_up_to), 'color': '#dc3545'})
                if hasattr(i, 'date_granted') and i.date_granted:
                    schedule.append({'title': f"{name} (GRANTED)", 'start': str(i.date_granted), 'color': '#1cc88a'})
        
        return render_template('deadlines.html', schedule=schedule)
        
    except Exception as e:
        import traceback
        return f"<div style='padding:20px; font-family:monospace; color:red;'><h2>CRASH REPORT:</h2><pre>{traceback.format_exc()}</pre></div>"

@app.route('/ai_insights/<asset_type>/<int:id>')
@login_required
def ai_insights(asset_type, id):
    return jsonify({"report": "AI Insight generated successfully."})

@app.route('/download_cert/<asset_type>/<int:id>')
@login_required
def download_cert(asset_type, id):
    item = None
    title, asset_id = "" , ""
    
    if asset_type == 'patent':
        item = PatentProduct.query.get(id)
        title, asset_id = item.patent_name, item.ip_asset_id
    elif asset_type == 'variety':
        item = Variety.query.get(id)
        title, asset_id = f"{item.category} (Released {item.year_of_release})", item.ip_asset_id
    elif asset_type == 'copyright':
        item = Copyright.query.get(id)
        title, asset_id = item.article_name, item.ip_asset_id
    elif asset_type == 'design':
        item = PatentDesign.query.get(id)
        title, asset_id = item.design_name, item.ip_asset_id

    if not item: return redirect(url_for('repository'))

    doc = Document()
    doc.add_heading('ICAR-IIRR IP Portfolio Summary', 0)
    doc.add_heading(f'Official {asset_type.capitalize()} Record', level=1)
    doc.add_paragraph(f"Registration ID: {asset_id}\nTitle/Details: {title}\nStatus: {item.status}")
    
    memory_file = io.BytesIO()
    doc.save(memory_file)
    memory_file.seek(0)
    return send_file(memory_file, as_attachment=True, download_name=f"{asset_id}_Certificate.docx")

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        hashed_pw = generate_password_hash(password)
        
        # New users are 'Scientist' by default and approved by Admin later
        new_user = User(username=username, email=email, password=hashed_pw, role='Scientist', is_approved=False)
        db.session.add(new_user)
        db.session.commit()
        
        flash('Access request submitted to Admin!')
        return redirect(url_for('login'))
    return render_template('signup.html')

@app.route('/manage_users')
@login_required
def manage_users():
    if current_user.role != 'Admin': return redirect(url_for('dashboard'))
    pending = User.query.filter_by(is_approved=False).all()
    active = User.query.filter_by(is_approved=True).all()
    return render_template('manage_users.html', pending=pending, active=active)

@app.route('/approve_user/<int:user_id>')
@login_required
def approve_user(user_id):
    # Security: only admins can approve users
    if current_user.role != 'Admin':
        return redirect(url_for('dashboard'))
    user = User.query.get(user_id)
    if user:
        # FIX: If you or admin sign up, you get Admin role automatically on approval
        if user.username.lower() in ['admin']:
            user.role = 'Admin'
        user.is_approved = True
        db.session.commit()
        
        try:
            msg = Message("ICAR-IIRR Account Approved", sender="suhasnethi04@gmail.com", recipients=[user.email])
            msg.body = f"Hello {user.username},\n\nYour account has been approved by the Admin.\n\nLink: http://127.0.0.1:5000/login"
            mail.send(msg)
            flash(f'User {user.username} approved!')
        except Exception as e:
            print(e)
            flash(f'User approved (Email failed).')

    return redirect(url_for('manage_users'))

@app.route('/api/suggest_assets')
@login_required
def suggest_assets():
    query = request.args.get('q', '').lower()
    category = request.args.get('cat', '')
    results = []
    
    if category in ['Variety', 'Hybrid']:
        items = Variety.query.filter(Variety.ip_asset_id.ilike(f'%{query}%')).all()
    elif category == 'Product':
        items = PatentProduct.query.filter(PatentProduct.ip_asset_id.ilike(f'%{query}%')).all()
    elif category == 'Process':
        items = PatentProcess.query.filter(PatentProcess.ip_asset_id.ilike(f'%{query}%')).all()
    elif category == 'Design':
        items = PatentDesign.query.filter(PatentDesign.ip_asset_id.ilike(f'%{query}%')).all()
    elif category == 'Copyright':
        items = Copyright.query.filter(Copyright.ip_asset_id.ilike(f'%{query}%')).all()
    elif category in ['Trademark', 'Logo']:
        items = Trademark.query.filter(Trademark.ip_asset_id.ilike(f'%{query}%')).all()
    else:
        return jsonify([])

    return jsonify([i.ip_asset_id for i in items])

@app.route('/import/<repo_type>', methods=['POST'])
@login_required
def import_csv_data(repo_type):
    if current_user.role != 'Admin': 
        return redirect(url_for('dashboard'))
    
    file = request.files.get('csv_file')
    if not file or file.filename == '': 
        flash("No file selected.")
        return redirect(url_for('repository_view', repo_type=repo_type))

    try:
        stream = io.StringIO(file.stream.read().decode("UTF8"), newline=None)
        reader = csv.DictReader(stream)
        
        if 'ID' not in reader.fieldnames:
            flash("Invalid CSV! Must contain an 'ID' column.")
            return redirect(url_for('repository_view', repo_type=repo_type))

        def p_date(d): 
            if not d or d.strip() in ['-', 'None', '']: return None
            try: return datetime.strptime(d.strip(), '%Y-%m-%d').date()
            except: return None

        def p_float(val):
            try: return float(str(val).replace(',', ''))
            except: return 0.0
            
        def p_int(val):
            try: return int(val)
            except: return None

        if repo_type == 'varieties':
            for row in reader:
                if Variety.query.filter_by(ip_asset_id=row['ID']).first(): continue
                def parse_json(v): return json.dumps([x.strip() for x in v.split(',')]) if v else '[]'
                v = Variety(
                    ip_asset_id=row['ID'], category=row.get('Cat'), status=row.get('Status', 'Filed'),
                    year_of_release=p_int(row.get('Year')), duration_days=p_int(row.get('Duration')),
                    ppvfr_registration=row.get('PPVFR'), agency_released_by=row.get('Notified By'),
                    yield_data=row.get('Yield'), recommended_zones=parse_json(row.get('Zones')),
                    states=parse_json(row.get('States')), grain_types=parse_json(row.get('Grain Types')), 
                    pest_resistances=parse_json(row.get('Pests')), disease_resistances=parse_json(row.get('Diseases')),
                    abiotic_stress=parse_json(row.get('Abiotic Stress')), special_traits=row.get('Special Traits')
                )
                db.session.add(v)

        elif repo_type == 'patents':
            for row in reader:
                ptype = row.get('Type', 'PatentProduct')
                model = PatentProduct if ptype == 'PatentProduct' else PatentProcess if ptype == 'PatentProcess' else PatentDesign
                if model.query.filter_by(ip_asset_id=row['ID']).first(): continue
                data = {
                    'ip_asset_id': row['ID'], 'status': row.get('Status', 'Filed'),
                    'date_filed': p_date(row.get('Date Filed')), 'date_granted': p_date(row.get('Date Granted')),
                    'valid_up_to': p_date(row.get('Expiry')),
                    'company_licensed': row.get('Company'), 'date_licensed': p_date(row.get('License / MoU Date')),
                    'license_fee': p_float(row.get('Fee')), 'license_fee_date': p_date(row.get('Fee Date')),
                    'royalty_received': p_float(row.get('Royalty')), 'date_royalty_received': p_date(row.get('Royalty Date'))
                }
                title = row.get('Title', 'Unnamed Asset')
                if ptype == 'PatentProduct': db.session.add(PatentProduct(patent_name=title, **data))
                elif ptype == 'PatentProcess': db.session.add(PatentProcess(process_name=title, **data))
                else: db.session.add(PatentDesign(design_name=title, **data))

        # --- BRANDS LOGIC (The missing piece!) ---
        elif repo_type == 'brands':
            for row in reader:
                ptype = row.get('Type', 'Copyright')
                if ptype == 'Copyright':
                    if Copyright.query.filter_by(ip_asset_id=row['ID']).first(): continue
                    new_item = Copyright(
                        ip_asset_id=row['ID'], article_name=row.get('Name', 'Unnamed'),
                        status=row.get('Status', 'Filed'), company_licensed=row.get('Company'),
                        date_licensed=p_date(row.get('License / MoU Date')),
                        license_fee=p_float(row.get('Fee')),
                        royalty_received=p_float(row.get('Royalty')),
                        date_royalty_received=p_date(row.get('Royalty Date'))
                    )
                else:
                    if Trademark.query.filter_by(ip_asset_id=row['ID']).first(): continue
                    new_item = Trademark(
                        ip_asset_id=row['ID'], asset_name=row.get('Name', 'Unnamed'),
                        asset_type=ptype, status=row.get('Status', 'Filed'),
                        company_licensed=row.get('Company'),
                        date_licensed=p_date(row.get('License / MoU Date')),
                        license_fee=p_float(row.get('Fee')),
                        royalty_received=p_float(row.get('Royalty')),
                        date_royalty_received=p_date(row.get('Royalty Date'))
                    )
                db.session.add(new_item)

        db.session.commit()
        flash("Data imported successfully!")
    except Exception as e:
        db.session.rollback()
        flash(f"Import Error: {str(e)}")

    return redirect(url_for('repository_view', repo_type=repo_type))

@app.route('/api/chat', methods=['POST'])
@login_required
def ai_chat():
    import re # <--- THIS IS REQUIRED FOR THE FIX
    user_msg = request.json.get('message', '').lower().strip()
    if not user_msg:
        return jsonify({'error': 'No message provided'}), 400

    # --- STEP 1: HARD FILTER ---
    greetings = ['hi', 'hello', 'hey', 'hii', 'help', 'who are you']
    if user_msg in greetings:
        return jsonify({'response': "Hello! I'm the ICAR-IIRR AI. Ask me about crop varieties or patents."})

    try:
        inspector = inspect(db.engine)
        schema_info = ""
        for table_name in inspector.get_table_names():
            if table_name != 'user': 
                cols = [col['name'] for col in inspector.get_columns(table_name)]
                schema_info += f"{table_name}({', '.join(cols)}); "

        # --- STEP 2: AI GENERATES SQL ---
        prompt1 = f"""Schema: {schema_info}
        User: {user_msg}
        Task: Write ONLY a valid SQLite SELECT query. 
        CRITICAL RULES:
        1. The name of the variety is ALWAYS stored in the 'ip_asset_id' column.
        2. 'yield_data' is a text string. To do math (like > 8.0), you MUST cast it: (CAST(yield_data AS REAL) >= 8.0) OR (yield_data LIKE '%8.%').
        3. The following columns are JSON arrays and MUST be searched using LIKE '%word%': recommended_zones, states, grain_types, pest_resistances, disease_resistances, abiotic_stress. Do NOT use '=' for these.
        If not possible, say 'NA'."""
        
        sql_res = client.models.generate_content(model='gemini-3-flash-preview', contents=prompt1)
        
        # --- BULLETPROOF SQL CLEANUP ---
        # This strips out ANY markdown tags (```sql, ```sqlite, ```json, etc.)
        query = re.sub(r'```[a-zA-Z]*', '', sql_res.text).replace('```', '').strip()
        
        print(f"\n--- DEBUG: AI GENERATED SQL ---\n{query}\n-------------------------------\n")

        if "SELECT" in query.upper():
            try:
                # --- SECURITY: Strict SQL validation + read-only execution ---
                query_clean = query.strip().rstrip(';')
                if ';' in query_clean or not query_clean.upper().startswith('SELECT'):
                    return jsonify({'response': "I can only run single search queries. Please rephrase."})
                
                with db.engine.connect() as ro_conn:
                    ro_conn.execute(db.text("PRAGMA query_only = ON"))
                    records = ro_conn.execute(db.text(query_clean)).fetchall()
                if not records:
                    return jsonify({'response': "I couldn't find any data for that in our current records."})
                
                clean_results = [" | ".join([str(val) for val in row if val is not None]) for row in records]
                
                # Protect the AI from massive data dumps to save quota
                data_dump = "\n".join(clean_results[:30]) 
                if len(clean_results) > 30:
                    data_dump += f"\n...and {len(clean_results) - 30} more records."

                # --- STEP 3: STRICT DATA ANSWER ---
                prompt2 = f"""You are a data assistant for ICAR-IIRR. 
                User asked: '{user_msg}'
                Raw database data:
                {data_dump}
                
                Task: Answer the user's question using ONLY the provided data. Give the exact answer with zero extra words, no introductory sentences, and no concluding notes."""
                final_res = client.models.generate_content(model='gemini-3-flash-preview', contents=prompt2)
                
                # Format for web display
                final_text = final_res.text.replace('\n', '<br>')
                return jsonify({'response': final_text})
                
            except Exception as db_err:
                print(f"\n--- DEBUG: DATABASE CRASH REASON ---\n{str(db_err)}\n------------------------------------\n")
                return jsonify({'response': "I had trouble reading the database for that specific question. Try rephrasing."})
        
        return jsonify({'response': "I can only help with database searches. Try asking about crops or patents."})

    except Exception as e:
        error_msg = str(e)
        print(f"API ERROR: {error_msg}")
        if "503" in error_msg:
            return jsonify({'response': "Google's AI servers are busy. Please click send again in 5 seconds."})
        if "429" in error_msg:
            return jsonify({'response': "The free tier is at its speed limit. Please wait 1 minute."})
        return jsonify({'response': "Error connecting to the AI service."})
    
if __name__ == '__main__':
    app.run(debug=True)